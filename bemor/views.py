from rest_framework.permissions import IsAdminUser
from rest_framework.parsers import MultiPartParser, FormParser
from django.http import HttpResponse
from rest_framework import generics, filters
from docx import Document
from .models import Bemor
from .serializers import BemorSerializer


class BemorCreateView(generics.CreateAPIView):
    queryset = Bemor.objects.all()
    serializer_class = BemorSerializer
    permission_classes = [IsAdminUser]


    def create(self, request, *args, **kwargs):
        response = super().create(request, *args, **kwargs)
        return response

class BemorRetrieveView(generics.RetrieveAPIView):
    queryset = Bemor.objects.all()
    serializer_class = BemorSerializer
    lookup_field = 'uid'

    def get(self, request, *args, **kwargs):
        bemor = self.get_object()  # Retrieve the Bemor object


        doc = Document()
        doc.add_heading('Bemor Ma\'lumotlari', level=1)

        # doc.add_paragraph(f"ID: {bemor.uid}")
        doc.add_paragraph(f"Ism: {bemor.name}")
        doc.add_paragraph(f"Kasallik: {bemor.kasallik}")
        doc.add_paragraph(f"Tashxis: {bemor.tashxis}")
        doc.add_paragraph(f"Shifokor: {bemor.doctor}")

        # Save to a temporary file
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=Bemor_{bemor.uid}.docx'
        doc.save(response)
        return response

class BemorUpdateView(generics.UpdateAPIView):
    queryset = Bemor.objects.all()
    serializer_class = BemorSerializer
    permission_classes = [IsAdminUser]
    lookup_field = 'uid'
    parser_classes = [MultiPartParser, FormParser]

class BemorDeleteView(generics.DestroyAPIView):
    queryset = Bemor.objects.all()
    serializer_class = BemorSerializer
    lookup_field = 'uid'
    permission_classes = [IsAdminUser]

class BemorListView(generics.ListAPIView):
    queryset = Bemor.objects.all().order_by('-uid')
    serializer_class = BemorSerializer
    permission_classes = [IsAdminUser]



class BemorDocxDownloadView(generics.RetrieveAPIView):
    queryset = Bemor.objects.all()
    serializer_class = BemorSerializer
    lookup_field = 'uid'
    permission_classes = []  # Open for everyone (no authentication required)

    def retrieve(self, request, *args, **kwargs):
        bemor = self.get_object()  # Get the specific Bemor instance

        # Create a new Word document
        doc = Document()
        doc.add_heading("Bemor Ma'lumotlari", level=1)

        # doc.add_paragraph(f"🆔 ID: {bemor.uid}")
        doc.add_paragraph(f" Ism: {bemor.name}")
        doc.add_paragraph(f"Kasallik: {bemor.kasallik}")
        doc.add_paragraph(f" Tashxis: {bemor.tashxis}")
        doc.add_paragraph(f"Shifokor: {bemor.doctor}")

        # Prepare the HTTP response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=Bemor_{bemor.uid}.docx'

        # Save the document to response
        doc.save(response)
        return response



class BemorSearchView(generics.ListAPIView):
    queryset = Bemor.objects.all()
    serializer_class = BemorSerializer
    permission_classes = []  # Open for everyone
    filter_backends = [filters.SearchFilter]
    search_fields = ['name']
